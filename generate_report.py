import os
import sys
import subprocess
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_border(cell, border_color_hex="D3D3D3"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
        border = OxmlElement(border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    set_cell_background(cell, "F5F5F5")
    set_cell_border(cell, "D3D3D3")
    
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_output_block(doc, output_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    set_cell_background(cell, "F0F4F8")  
    set_cell_border(cell, "B0C4DE")
    
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(output_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x24, 0x21)

def read_file_safe(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    return f"# Помилка: Файл {os.path.basename(file_path)} не знайдено."

def main():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_top = p_top.add_run("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ\nФАХОВИЙ КОЛЕДЖ")
    run_top.font.name = 'Times New Roman'
    run_top.font.size = Pt(12)
    run_top.font.bold = True
    p_top.paragraph_format.space_after = Pt(100)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ЗВІТ\nПРО ВИКОНАННЯ ЛАБОРАТОРНОЇ РОБОТИ №1")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    p_title.paragraph_format.space_after = Pt(12)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("з дисципліни «Програмування та основи алгоритмізації»\n"
                            "на тему: «Основи алгоритмізації та базові структури даних на мові Python»")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    p_sub.paragraph_format.space_after = Pt(120)

    p_stud = doc.add_paragraph()
    p_stud.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_stud = p_stud.add_run(
        "Виконав:\n"
        "Студент групи КІТ-42\n"
        "Павлченко Платон\n\n"
        "Репозиторій: https://github.com/plton228/Algorithmization.git"
    )
    run_stud.font.name = 'Times New Roman'
    run_stud.font.size = Pt(12)
    p_stud.paragraph_format.space_after = Pt(100)

    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_bot = p_bot.add_run("2026")
    run_bot.font.name = 'Times New Roman'
    run_bot.font.size = Pt(12)
    
    doc.add_page_break()

    p_meta_title = doc.add_paragraph()
    run_meta_title = p_meta_title.add_run("МЕТА РОБОТИ")
    run_meta_title.font.name = 'Times New Roman'
    run_meta_title.font.size = Pt(14)
    run_meta_title.font.bold = True
    p_meta_title.paragraph_format.space_before = Pt(12)
    p_meta_title.paragraph_format.space_after = Pt(6)

    p_meta_text = doc.add_paragraph()
    run_meta_text = p_meta_text.add_run(
        "Набути початкових практичних навичок реалізації простих алгоритмів мовою програмування Python; "
        "закріпити вміння працювати з циклами, умовами, списками (масивами) та рядками; "
        "ознайомитися з базовими структурами даних - стеком і чергою; "
        "навчитися застосовувати прості алгоритми пошуку та сортування й перевіряти "
        "правильність роботи програм на контрольних прикладах."
    )
    run_meta_text.font.name = 'Times New Roman'
    run_meta_text.font.size = Pt(12)
    p_meta_text.paragraph_format.line_spacing = 1.15
    p_meta_text.paragraph_format.space_after = Pt(18)

    p_repo = doc.add_paragraph()
    run_repo_lbl = p_repo.add_run("Репозиторій GitHub: ")
    run_repo_lbl.font.name = 'Times New Roman'
    run_repo_lbl.font.size = Pt(12)
    run_repo_lbl.font.bold = True
    run_repo_url = p_repo.add_run("https://github.com/plton228/Algorithmization.git")
    run_repo_url.font.name = 'Times New Roman'
    run_repo_url.font.size = Pt(12)
    run_repo_url.font.underline = True
    run_repo_url.font.color.rgb = RGBColor(0x00, 0x00, 0xEE)
    p_repo.paragraph_format.space_after = Pt(18)

    tasks_info = [
        {
            "num": 1,
            "title": "Базові конструкції: цикли та умови",
            "condition": "Написати програму, яка приймає від користувача натуральне число N і виводить усі парні числа від 1 до N, а також обчислює їхню суму. Програма має використовувати цикл та умовний оператор для перевірки парності числа.",
            "desc": "Алгоритм виконує ітерацію за допомогою циклу `for` по всіх цілих числах від 1 до N включно. На кожному кроці за допомогою умовного оператора `if` і операції взяття остачі від ділення `% 2` перевіряється парність поточного числа. Якщо воно парне - додається до списку та загальної суми.",
            "file": "task1_loops_conditions.py",
            "demo_output": (
                "--- Завдання 1: Парні числа та їх сума від 1 до N ---\n"
                "Введіть натуральне число N (N >= 1): 15\n"
                "Парні числа в діапазоні від 1 до 15: 2, 4, 6, 8, 10, 12, 14\n"
                "Сума парних чисел: 56"
            )
        },
        {
            "num": 2,
            "title": "Робота з рядками",
            "condition": "Написати програму, яка приймає від користувача рядок тексту й підраховує кількість голосних та приголосних літер у ньому. Великі та малі літери вважати однаковими. Додатково вивести довжину рядка.",
            "desc": "Рядок перетворюється на нижній регістр за допомогою методу `.lower()`. Далі створюються дві множини (`set`) символів: одна містить голосні літери (українські та англійські), а інша - приголосні літери. Програма перебирає кожен символ у рядку та перевіряє його належність до відповідних множин, збільшуючи лічильники. Довжина рядка визначається за допомогою вбудованої функції `len()`.",
            "file": "task2_strings.py",
            "demo_output": (
                "--- Завдання 2: Аналіз рядка тексту ---\n"
                "Введіть рядок тексту: Привіт, Python 2026!\n"
                "Загальна довжина рядка: 20 символів\n"
                "Кількість голосних літер: 4\n"
                "Кількість приголосних літер: 9"
            )
        },
        {
            "num": 3,
            "title": "Робота зі списками (масивами)",
            "condition": "Написати програму, яка створює список із N цілих чисел, введених користувачем, і знаходить у ньому найбільший та найменший елементи, а також обчислює середнє арифметичне всіх чисел. Пошук максимуму й мінімуму виконати самостійно за допомогою циклу (без використання вбудованих функцій max і min).",
            "desc": "Програма зчитує кількість елементів N та заповнює список значеннями, які користувач вводить вручну. Перший елемент списку приймається за початкові значення мінімуму (`min_val`) та максимуму (`max_val`). Під час проходу по списку у циклі `for` кожен наступний елемент порівнюється з поточними екстремумами, оновлюючи their значення у разі потреби. Паралельно підраховується сума елементів, яка ділиться на N для знаходження середнього арифметичного.",
            "file": "task3_lists.py",
            "demo_output": (
                "--- Завдання 3: Пошук мінімуму, максимуму та середнього у списку ---\n"
                "Введіть кількість чисел N: 5\n"
                "Введіть число 1: 12\n"
                "Введіть число 2: -5\n"
                "Введіть число 3: 8\n"
                "Введіть число 4: 25\n"
                "Введіть число 5: 3\n"
                "Створений список чисел: [12, -5, 8, 25, 3]\n"
                "Найменший елемент: -5\n"
                "Найбільший елемент: 25\n"
                "Середнє арифметичне: 8.60"
            )
        },
        {
            "num": 4,
            "title": "Лінійний пошук",
            "condition": "Реалізувати алгоритм лінійного пошуку: програма отримує список чисел і шукане число, після чого по черзі переглядає елементи списку та повідомляє, чи є таке число у списку, і на якій позиції (індексі) воно знаходиться. Продемонструвати роботу на прикладі, коли число є у списку, і коли його немає.",
            "desc": "Лінійний пошук послідовно перевіряє кожен елемент списку за допомогою циклу `for` за його індексом. Якщо поточний елемент дорівнює шуканому значенню, функція негайно повертає поточний індекс. Якщо цикл завершується без виявлення збігів, повертається значення -1, що сигналізує про відсутність елемента в списку.",
            "file": "task4_linear_search.py",
            "demo_output": (
                "--- Демонстрація роботи лінійного пошуку ---\n"
                "Початковий список для демонстрації: [10, 25, 3, 47, 8, 19, 4]\n\n"
                "Шукаємо число: 8\n"
                "Результат: Число 8 знайдено на позиції (індексі): 4\n\n"
                "Шукаємо число: 99\n"
                "Результат: Число 99 відсутнє у списку\n"
                "---------------------------------------------\n"
                "Бажаєте спробувати самостійно?\n"
                "Введіть список чисел через пробіл (або натисніть Enter для пропуску): "
            )
        },
        {
            "num": 5,
            "title": "Бульбашкове сортування",
            "condition": "Реалізувати алгоритм бульбашкового сортування для впорядкування списку чисел за зростанням. Вивести список до та після сортування. Заборонено використовувати вбудовані функції sorted або sort — алгоритм потрібно написати самостійно за допомогою вкладених циклів.",
            "desc": "Алгоритм бульбашкового сортування використовує вкладені цикли `for`. Зовнішній цикл визначає кількість проходів по масиву, а внутрішній порівнює сусідні елементи. Якщо попередній елемент більший за наступний, вони міняються місцями (swapping). Для оптимізації використовується прапорець `swapped`, який дозволяє достроково завершити сортування, якщо під час проходу не було зроблено жодного обміну (масив уже відсортовано).",
            "file": "task5_bubble_sort.py",
            "demo_output": (
                "--- Демонстрація роботи бульбашкового сортування ---\n"
                "Список до сортування : [64, 34, 25, 12, 22, 11, 90]\n"
                "Список після сортування: [11, 12, 22, 25, 34, 64, 90]\n"
                "----------------------------------------------------\n"
                "Бажаєте спробувати самостійно?\n"
                "Введіть список чисел через пробіл (або натисніть Enter для пропуску): "
            )
        },
        {
            "num": 6,
            "title": "Бінарний пошук",
            "condition": "Реалізувати алгоритм бінарного (двійкового) пошуку у вже відсортованому за зростанням списку чисел. Програма отримує відсортований список і шукане число та знаходить його позицію, щоразу ділячи область пошуку навпіл. Продемонструвати роботу на прикладі, коли елемент є у списку, і коли його немає.",
            "desc": "Алгоритм вимагає відсортованого списку. Він ініціалізує ліву (`low` = 0) та праву (`high` = довжина списку - 1) межі пошуку. У циклі знаходить середній індекс `mid = (low + high) // 2`. Якщо шуканий елемент дорівнює елементу на позиції `mid`, пошук успішно завершено. Якщо шукане число більше за елемент `mid`, ліва межа переміщується на `mid + 1`. Якщо менше - права межа переміщується на `mid - 1`. Цикл працює, поки `low <= high`, повертаючи `-1` при невдачі.",
            "file": "task6_binary_search.py",
            "demo_output": (
                "--- Демонстрація роботи бінарного пошуку ---\n"
                "Відсортований список для демонстрації: [3, 8, 12, 19, 25, 34, 47, 64, 90]\n\n"
                "Шукаємо число: 25\n"
                "Результат: Число 25 знайдено на індексі: 4\n\n"
                "Шукаємо число: 50\n"
                "Результат: Число 50 відсутнє у списку\n"
                "--------------------------------------------------\n"
                "Бажаєте спробувати самостійно?\n"
                "Введіть список чисел через пробіл (вони будуть автоматично відсортовані): "
            )
        },
        {
            "num": 7,
            "title": "Структура даних: стек",
            "condition": "Реалізувати просту структуру даних «стек» на основі списку Python з операціями: push (додати елемент), pop (вилучити верхній елемент), peek (переглянути верхній елемент) та is_empty (перевірити, чи стек порожній). Продемонструвати роботу стеку: додати кілька елементів, переглянути верхній, вилучити кілька елементів і вивести стан стеку після кожної операції.",
            "desc": "Стек реалізовано як об'єктно-орієнтований клас `Stack` на базі стандартного списку. Метод `push` додає елемент у кінець списку (`.append()`), `pop` вилучає та повертає останній елемент (`.pop()`), `peek` повертає останній елемент без вилучення (`self.items[-1]`), а `is_empty` перевіряє довжину списку. Це реалізує принцип LIFO (Last In, First Out) зі складністю O(1) для основних операцій.",
            "file": "task7_stack.py",
            "demo_output": (
                "--- Завдання 7: Структура даних Стек (LIFO) ---\n"
                "Ініціалізовано новий стек. Чи порожній він? True\n"
                "Стек (дно -> вершина): []\n"
                "--------------------------------------------------\n"
                "Послідовно додаємо елементи: [10, 20, 30, 40]\n"
                "Додано: 10 | Стек (дно -> вершина): [10]\n"
                "Додано: 20 | Стек (дно -> вершина): [10, 20]\n"
                "Додано: 30 | Стек (дно -> вершина): [10, 20, 30]\n"
                "Додано: 40 | Стек (дно -> вершина): [10, 20, 30, 40]\n"
                "--------------------------------------------------\n"
                "Перегляд вершини стеку (peek): 40\n"
                "Стек (дно -> вершина): [10, 20, 30, 40]\n"
                "--------------------------------------------------\n"
                "Вилучаємо елементи зі стеку:\n"
                "Вилучено: 40 | Стек (дно -> вершина): [10, 20, 30]\n"
                "Вилучено: 30 | Стек (дно -> вершина): [10, 20]\n"
                "Вилучено: 20 | Стек (дно -> вершина): [10]\n"
                "Вилучено: 10 | Стек (дно -> вершина): []\n"
                "--------------------------------------------------\n"
                "Спроба вилучити елемент з уже порожнього стеку:\n"
                "Помилка: Спроба вилучення елемента з порожнього стеку!\n"
                "Стек (дно -> вершина): []"
            )
        }
    ]

    for t in tasks_info:
        p_task_title = doc.add_paragraph()
        run_task_title = p_task_title.add_run(f"Завдання {t['num']}. {t['title']}")
        run_task_title.font.name = 'Times New Roman'
        run_task_title.font.size = Pt(13)
        run_task_title.font.bold = True
        p_task_title.paragraph_format.space_before = Pt(18)
        p_task_title.paragraph_format.space_after = Pt(6)

        p_cond = doc.add_paragraph()
        p_cond.paragraph_format.left_indent = Inches(0.25)
        run_cond_lbl = p_cond.add_run("Умова задачі: ")
        run_cond_lbl.font.name = 'Times New Roman'
        run_cond_lbl.font.size = Pt(11)
        run_cond_lbl.font.bold = True
        run_cond_txt = p_cond.add_run(t['condition'])
        run_cond_txt.font.name = 'Times New Roman'
        run_cond_txt.font.size = Pt(11)
        p_cond.paragraph_format.space_after = Pt(4)

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent = Inches(0.25)
        run_desc_lbl = p_desc.add_run("Короткий опис алгоритму: ")
        run_desc_lbl.font.name = 'Times New Roman'
        run_desc_lbl.font.size = Pt(11)
        run_desc_lbl.font.bold = True
        run_desc_txt = p_desc.add_run(t['desc'])
        run_desc_txt.font.name = 'Times New Roman'
        run_desc_txt.font.size = Pt(11)
        p_desc.paragraph_format.space_after = Pt(6)

        code_content = read_file_safe(t['file'])
        p_code_lbl = doc.add_paragraph()
        run_code_lbl = p_code_lbl.add_run("Вихідний код програми:")
        run_code_lbl.font.name = 'Times New Roman'
        run_code_lbl.font.size = Pt(11)
        run_code_lbl.font.bold = True
        p_code_lbl.paragraph_format.space_after = Pt(2)
        add_code_block(doc, code_content)
        
        p_out_lbl = doc.add_paragraph()
        run_out_lbl = p_out_lbl.add_run("Приклад введення та виведення програми:")
        run_out_lbl.font.name = 'Times New Roman'
        run_out_lbl.font.size = Pt(11)
        run_out_lbl.font.bold = True
        p_out_lbl.paragraph_format.space_before = Pt(6)
        p_out_lbl.paragraph_format.space_after = Pt(2)
        add_output_block(doc, t['demo_output'])
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(12)

    doc.add_page_break()
    p_conc_title = doc.add_paragraph()
    run_conc_title = p_conc_title.add_run("ВИСНОВКИ")
    run_conc_title.font.name = 'Times New Roman'
    run_conc_title.font.size = Pt(14)
    run_conc_title.font.bold = True
    p_conc_title.paragraph_format.space_before = Pt(12)
    p_conc_title.paragraph_format.space_after = Pt(6)

    p_conc_text = doc.add_paragraph()
    run_conc_text = p_conc_text.add_run(
        "Під час виконання лабораторної роботи №1 було детально вивчено та практично застосовано основні синтаксичні конструкції мови програмування Python. "
        "Зокрема, закріплено роботу з умовними операторами та різними видами циклів (for та while).\n\n"
        "Набуто практичних навичок маніпулювання списками (масивами) та текстовими рядками, зокрема підрахунку різних груп символів у тексті та самостійного пошуку екстремумів без вбудованих функцій.\n\n"
        "Практично реалізовано базові алгоритми лінійного та бінарного пошуків, а також сортування бульбашкою, що дозволило краще зрозуміти логіку роботи з індексами та ефективність алгоритмів.\n\n"
        "Також було реалізовано структуру даних «стек» за принципом LIFO, що продемонструвало переваги використання об'єктно-орієнтованого підходу в Python для структурування даних."
    )
    run_conc_text.font.name = 'Times New Roman'
    run_conc_text.font.size = Pt(12)
    p_conc_text.paragraph_format.line_spacing = 1.15
    p_conc_text.paragraph_format.space_after = Pt(12)

    output_filename = "Звіт_Лабораторна_1.docx"
    try:
        doc.save(output_filename)
        print(f"Report successfully generated and saved to {output_filename}")
    except PermissionError:
        fallback_filename = "Звіт_Лабораторна_1_Платон.docx"
        doc.save(fallback_filename)
        print(f"Permission denied for {output_filename} (probably open in Word). Saved report as {fallback_filename} instead.")

if __name__ == "__main__":
    main()
